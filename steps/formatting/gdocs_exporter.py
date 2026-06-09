"""
Google Docs exporter — renders merge_remediation.json as a reviewable .docx file.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

_GRAY = (110, 110, 110)

_TOOL_LABELS = {
    "multiple_choice": "Select the correct answer",
    "multi_select": "Select ALL correct answers",
    "click_component": "Click the correct component",
    "click_category": "Click on a category",
    "click_tangible": "Click on an object",
    "click_to_place": "Click to place",
    "click_place_symbols": "Place symbols on the graph",
    "click_to_set_height": "Set the bar height",
    "click_scale_button": "Set the scale",
    "drag_to_sort": "Drag to sort",
    "place_tile": "Place a tile",
    "add_item_per_row": "Add items to each row",
    "add_row": "Add a row",
    "set_container_count": "Set the number of containers",
    "set_items_per_container": "Set items per container",
    "counting_game": "Count",
    "data_collection_interaction": "Collect data",
}

_GLOSSARY = [
    ("Guide:", "The tutor's spoken dialogue"),
    ("Visual:", "What appears on screen (visual setup, animations, highlights)"),
    ("Student Interaction:", "The type of action the student performs"),
    ("Prompt:", "The specific question or instruction the student receives"),
    ("• Answer option (optional):", "The available answer choice, nested under Prompt"),
    ("🟢  Correct", "Student answers correctly, student moves forward"),
    ("🟡  Attempt 1", "Student answers wrong on first attempt — light remediation"),
    ("🟠  Attempt 2", "Student answers wrong on second attempt — medium remediation"),
    (
        "🟠  Pattern branch",
        "Response triggered by a specific mistake pattern (not attempt-count-driven)",
    ),
    (
        "🔴  Heavy / Fallback",
        "Wrong 3+ times, system models the correct answer and the student moves forward",
    ),
]


# ── Formatting helpers ─────────────────────────────────────────────────────────


def _tight(para, before: int = 0, after: int = 2):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def _run(para, text: str, *, bold=False, italic=False, color=None, size=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = RGBColor(*color)
    if size:
        r.font.size = Pt(size)


# ── Starter pack discovery & parsing ──────────────────────────────────────────


def _discover_starter_pack(json_path: Path) -> tuple[Path | None, str | None]:
    """Return (starter_pack_md_path, phase_name) inferred from a tracked_scripts path."""
    parts = json_path.resolve().parts
    try:
        ts_idx = next(i for i, p in enumerate(parts) if p == "tracked_scripts")
    except StopIteration:
        return None, None

    if ts_idx + 3 >= len(parts):
        return None, None

    m_unit = re.match(r"u(\d+)", parts[ts_idx + 1])
    m_mod = re.match(r"m(\d+)", parts[ts_idx + 2])
    phase = parts[ts_idx + 3]
    if not m_unit or not m_mod:
        return None, None

    root = Path(*parts[:ts_idx])
    sp = (
        root
        / "units"
        / f"unit{m_unit.group(1)}"
        / f"module{m_mod.group(1)}"
        / "_starter_pack_ref.md"
    )
    return (sp if sp.exists() else None), phase


def _parse_starter_pack(md_path: Path) -> dict:
    text = md_path.read_text(encoding="utf-8")

    # Module title — first top-level heading, strip markdown bold markers
    title = ""
    m = re.search(r"^#\s+\*{0,2}(.+?)\*{0,2}\s*$", text, re.MULTILINE)
    if m:
        title = m.group(1).strip()

    # The One Thing — bold paragraph immediately after the §1.0 heading
    one_thing = ""
    m = re.search(r"THE ONE THING\s*\n+\*{2}(.+?)\*{2}", text, re.DOTALL)
    if m:
        one_thing = re.sub(r"\s+", " ", m.group(1)).strip()

    # Learning goals — **L1:** / **L2:** lines
    goals = []
    for gm in re.finditer(r"\*\*(L\d+):\*\*\s*(.+?)\s*$", text, re.MULTILINE):
        goal_text = gm.group(2).strip().strip('"').strip("“").strip("”")
        goals.append((gm.group(1), goal_text))

    # Misconceptions — ### sub-headings inside §1.4
    misconceptions = []
    section_14 = re.search(r"##\s+\*{0,2}1\.4 MISCONCEPTIONS.*?(?=\n##\s)", text, re.DOTALL)
    if section_14:
        for mm in re.finditer(r"###\s+\*{0,2}(.+?)\*{0,2}\s*$", section_14.group(), re.MULTILINE):
            label = mm.group(1).strip().replace(r"\#", "#")
            misconceptions.append(label)

    # Addressing standard — §1.1.1 Standards Cascade table, "Addressing" row
    addressing_standard = ""
    m = re.search(r"\|\s*\*{0,2}Addressing\*{0,2}\s*\|\s*(.+?)\s*\|", text, re.IGNORECASE)
    if m:
        addressing_standard = m.group(1).strip()

    # Module bridges — §1.1.2: "From ..." and "To Module N" bold paragraphs
    building_from = ""
    building_toward = ""
    bridges_section = re.search(r"1\.1\.2 Module Bridges(.*?)(?=\n#{2,3}\s|\Z)", text, re.DOTALL)
    if bridges_section:
        blob = bridges_section.group(1)
        m = re.search(r"\*\*From[^*]*\*\*[:\s]*(.*?)(?=\n\*\*|\Z)", blob, re.DOTALL)
        if m:
            building_from = re.sub(r"\s+", " ", m.group(1)).strip()
        m = re.search(r"\*\*To Module[^*]*\*\*[:\s]*(.*?)(?=\n\*\*|\Z)", blob, re.DOTALL)
        if m:
            building_toward = re.sub(r"\s+", " ", m.group(1)).strip()

    return {
        "title": title,
        "one_thing": one_thing,
        "goals": goals,
        "addressing_standard": addressing_standard,
        "misconceptions": misconceptions,
        "building_from": building_from,
        "building_toward": building_toward,
    }


# ── Header & glossary renderers ───────────────────────────────────────────────


def _render_module_context(doc: Document, module_info: dict, phase: str | None):
    # Module title
    h = doc.add_heading(module_info["title"] or "Module Review", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(2)

    if phase:
        p = doc.add_paragraph()
        _tight(p, before=0, after=6)
        _run(p, f"Phase: {phase.title()}", italic=True, color=_GRAY)

    if module_info["goals"]:
        p = doc.add_paragraph()
        _tight(p, before=4, after=2)
        _run(p, "Learning Goals", bold=True)
        for label, text in module_info["goals"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            _tight(p, before=0, after=1)
            _run(p, f"{label}: ", bold=True)
            _run(p, text)
        if module_info.get("addressing_standard"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            _tight(p, before=2, after=1)
            _run(p, "Common Core Standard Addressed: ", bold=True)
            _run(p, module_info["addressing_standard"])

    if module_info["one_thing"]:
        p = doc.add_paragraph()
        _tight(p, before=6, after=2)
        _run(p, "The One Thing: ", bold=True)
        _run(p, module_info["one_thing"])

    if module_info["misconceptions"]:
        p = doc.add_paragraph()
        _tight(p, before=6, after=2)
        _run(p, "Key Misconceptions", bold=True)
        for misc in module_info["misconceptions"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            _tight(p, before=0, after=1)
            _run(p, f"• {misc}")

    if module_info.get("building_from"):
        p = doc.add_paragraph()
        _tight(p, before=6, after=2)
        _run(p, "Building From: ", bold=True)
        _run(p, module_info["building_from"])

    if module_info.get("building_toward"):
        p = doc.add_paragraph()
        _tight(p, before=2, after=2)
        _run(p, "Building Toward: ", bold=True)
        _run(p, module_info["building_toward"])

    # Spacer before glossary
    p = doc.add_paragraph()
    _tight(p, before=8, after=0)


def _render_glossary(doc: Document):
    h = doc.add_heading("Document Guide", level=2)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(4)

    for label, desc in _GLOSSARY:
        p = doc.add_paragraph()
        _tight(p, before=1, after=1)
        _run(p, f"{label}  ", bold=True)
        _run(p, desc, color=_GRAY, size=9.5)

    p = doc.add_paragraph()
    _tight(p, before=10, after=0)


# ── Script data helpers ────────────────────────────────────────────────────────


def _prettify_id(section_id: str) -> str:
    m = re.match(r"s(\d+)_(\d+)_(.*)", section_id)
    if m:
        major, minor, rest = m.groups()
        name = rest.replace("_", " ").title()
        return f"S{major}.{minor} — {name}"
    return section_id.replace("_", " ").title()


def _scene_desc(beat: dict) -> str:
    method = beat.get("method", "add")
    params = beat.get("params", {})
    desc = params.get("description", "")
    if desc:
        return desc

    parts = []
    highlight = params.get("highlight_categories", [])
    dim = params.get("dim_categories", [])
    if highlight:
        parts.append(f"Highlight: {', '.join(highlight)}")
    if dim:
        parts.append(f"Dim: {', '.join(dim)}")
    if parts:
        return " / ".join(parts)

    tangible = beat.get("tangible_id", "")
    ttype = beat.get("tangible_type", "").replace("_", " ")
    if method == "remove":
        return f"Remove {ttype or tangible}"
    return " ".join(
        p for p in [method.capitalize(), ttype, f"({tangible})" if tangible else ""] if p
    )


def _flatten(container: dict) -> list[dict]:
    if "beats" in container:
        return container["beats"]
    if "steps" in container:
        result = []
        for group in container["steps"]:
            result.extend(group)
        return result
    return []


def _branch_desc(raw: str) -> str:
    return re.sub(r"^Branch\s+\d+:\s*", "", raw).strip()


def _find_incorrect_count(condition) -> int | None:
    if isinstance(condition, dict):
        if "incorrect_count" in condition:
            return int(condition["incorrect_count"])
        for v in condition.values():
            result = _find_incorrect_count(v)
            if result is not None:
                return result
    elif isinstance(condition, list):
        for item in condition:
            result = _find_incorrect_count(item)
            if result is not None:
                return result
    return None


def _classify(v: dict) -> tuple[str, str]:
    """Return (emoji, cleaned_description).

    🟢  correct
    🟡  incorrect_count == 1  (light / attempt 1)
    🟠  incorrect_count == 2  OR pattern-based (no attempt count)
    🔴  incorrect_count >= 3  OR empty condition (system-model fallback)
    """
    desc = _branch_desc(v.get("description", ""))
    if v.get("is_correct"):
        return "🟢", desc
    condition = v.get("condition", {})
    if not condition:
        return "🔴", desc
    count = _find_incorrect_count(condition)
    if count is None:
        return "🟠", desc
    if count == 1:
        return "🟡", desc
    if count == 2:
        return "🟠", desc
    return "🔴", desc


# ── Script renderers ───────────────────────────────────────────────────────────


def _render_guide(doc: Document, text: str, indent: float):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    _tight(p, before=1, after=1)
    _run(p, "Guide: ", bold=True)
    _run(p, text)


def _render_visual(doc: Document, text: str, indent: float):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    _tight(p, before=1, after=1)
    _run(p, "Visual: ", bold=True)
    _run(p, text)


def _render_prompt(doc: Document, beat: dict, indent: float):
    tool_raw = beat.get("tool", "")
    tool_label = _TOOL_LABELS.get(tool_raw, tool_raw.replace("_", " ").title() or "Unknown")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    _tight(p, before=4, after=1)
    _run(p, "Student Interaction: ", bold=True)
    _run(p, tool_label)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    _tight(p, before=1, after=1)
    _run(p, "Prompt: ", bold=True)
    _run(p, beat.get("text", ""), italic=True)

    for opt in beat.get("options", []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent + 0.25)
        _tight(p, before=0, after=1)
        _run(p, f"• {opt}")

    validator = beat.get("validator", [])
    if validator:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        _tight(p, before=4, after=2)
        _run(p, "Based on student response:", bold=True)

    for v in validator:
        emoji, desc = _classify(v)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        _tight(p, before=2, after=1)
        _run(p, f"{emoji}  ", bold=True)
        _run(p, desc, bold=True, size=9.5)

        for rb in _flatten(v):
            _render_beat(doc, rb, indent=indent)


def _render_beat(doc: Document, beat: dict, indent: float = 0.0):
    btype = beat.get("type")
    if btype == "dialogue":
        _render_guide(doc, beat.get("text", ""), indent)
    elif btype == "scene":
        desc = _scene_desc(beat)
        if desc:
            _render_visual(doc, desc, indent)
    elif btype == "prompt":
        _render_prompt(doc, beat, indent)
    # current_scene: skip


def _render_section(doc: Document, section: dict):
    h = doc.add_heading(_prettify_id(section.get("id", "unknown")), level=1)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    beats = _flatten(section)
    start = 0
    while (
        start < len(beats)
        and beats[start].get("type") == "scene"
        and beats[start].get("method") == "remove"
    ):
        start += 1
    for beat in beats[start:]:
        _render_beat(doc, beat)


# ── Entry points ──────────────────────────────────────────────────────────────

_PHASE_ORDER = ["warmup", "lesson", "exitcheck", "synthesis"]


def _find_phase_jsons(module_dir: Path) -> list[tuple[str, Path]]:
    """Return (phase, json_path) pairs for a module directory, in canonical phase order."""
    results = []
    for phase in _PHASE_ORDER:
        matches = sorted((module_dir / phase).rglob("merge_remediation.json"))
        if matches:
            results.append((phase, matches[0]))
    return results


def _find_json_in_phase(phase_dir: Path, kind: str) -> Path | None:
    """Find merge_remediation.json or pull.json in a phase directory."""
    if kind == "merge_remediation":
        matches = sorted(phase_dir.glob("step_*merge_remediation/merge_remediation.json"))
    else:
        matches = sorted(phase_dir.glob("step_*pull/pull.json"))
    return matches[0] if matches else None


def phase_export(phase_dir: Path, output_path: Path, kind: str) -> bool:
    """Generate a per-phase docx from merge_remediation or pull json. Returns True if written."""
    json_path = _find_json_in_phase(phase_dir, kind)
    if not json_path:
        return False
    sections = json.loads(json_path.read_text(encoding="utf-8"))
    to_docx(sections, output_path, json_path=json_path)
    print(f"  Saved: {output_path}")
    return True


def _module_number(module_dir: Path) -> str:
    """Extract module number from a path component like 'm4' → '4', 'm12' → '12'."""
    m = re.search(r"m(\d+)", module_dir.name)
    return m.group(1) if m else module_dir.name


def run_phase_exports(module_dir: Path):
    """Generate per-phase docx files for a module."""
    num = _module_number(module_dir)
    for phase in _PHASE_ORDER:
        phase_dir = module_dir / phase
        if not phase_dir.exists():
            continue
        phase_export(
            phase_dir, phase_dir / f"Module {num} - Unedited Script.docx", "merge_remediation"
        )
        phase_export(
            phase_dir,
            phase_dir / f"Module {num} - Teacher Reviewed Script (comparison reference).docx",
            "pull",
        )


def batch_export(module_dir: Path, output_path: Path):
    """Generate a single review docx combining all phases for a module."""
    phase_jsons = _find_phase_jsons(module_dir)
    if not phase_jsons:
        print(f"  No merge_remediation.json found in {module_dir}")
        return

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet", "List Paragraph"):
        try:
            doc.styles[style_name].font.name = "Calibri"
        except KeyError:
            pass

    sp_path, _ = _discover_starter_pack(phase_jsons[0][1])
    if sp_path:
        _render_module_context(doc, _parse_starter_pack(sp_path), phase=None)
    _render_glossary(doc)

    for phase, json_path in phase_jsons:
        h = doc.add_heading(phase.title(), level=1)
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(6)
        sections = json.loads(json_path.read_text(encoding="utf-8"))
        for section in sections:
            _render_section(doc, section)

    doc.save(output_path)
    print(f"  Saved: {output_path}")


def to_docx(sections: list[dict], output_path: Path, json_path: Path | None = None):
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet", "List Paragraph"):
        try:
            doc.styles[style_name].font.name = "Calibri"
        except KeyError:
            pass

    sp_path, phase = _discover_starter_pack(json_path) if json_path else (None, None)
    if sp_path:
        _render_module_context(doc, _parse_starter_pack(sp_path), phase)
    _render_glossary(doc)

    for section in sections:
        _render_section(doc, section)

    doc.save(output_path)


def export(data, output_path: str | None = None, **_kwargs):
    """Pipeline-compatible entry point. Passes data through unchanged."""
    sections = data if isinstance(data, list) else [data]
    path = Path(output_path or "review.docx")
    to_docx(sections, path)
    print(f"[gdocs_exporter] Saved {path}")
    return data


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python gdocs_exporter.py <input.json> [output.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    sections = json.loads(src.read_text(encoding="utf-8"))
    to_docx(sections, dst, json_path=src)
    print(f"Saved: {dst}")
